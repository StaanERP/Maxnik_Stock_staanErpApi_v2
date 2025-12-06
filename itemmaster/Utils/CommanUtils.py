from decimal import Decimal
import json
from django.core.exceptions import ValidationError
from userManagement.models import UserManagement
from ..models import *
from django.db.models import ProtectedError
import requests
import json
import base64
import os
from django.apps import apps
from django.core.exceptions import ObjectDoesNotExist
from docx import Document as Documentpdf
from docx.oxml import OxmlElement,parse_xml
from docx.oxml.ns import qn,nsdecls
from docx.shared import Pt,Inches
from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_HEADER_FOOTER
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from django.core.cache import cache
from django.db.models import Sum, Case, When, Value, IntegerField
from django.db.models.functions import Cast
from django.db.models import FloatField

def CreateOrEditContact(id=None, name=None, email=None, mobile_number=None, alternate_mobile_number=None, type=None):
    """
    Create And Upadet the contact
    """
    success = False
    errors = []
    data = None
    contact_type_ = None
    try:
        contact_type_ = Contact_type.objects.get(name=type)
    except Exception as e:
        errors.append(e)
    if contact_type_ is None:
        errors.append("Ask admin to add enquiry Type in Contact.")
        return {"data": data, "success": success, "errors": errors}
    if id:
        try:
            # Fetch the contact if the id is provided
            contact = ContactDetalis.objects.get(id=id)
            # Update fields if the new values are different from the existing ones
            contact.contact_person_name = name if name is not None else contact.contact_person_name
            contact.email = email if email is not None else contact.email
            contact.phone_number = mobile_number if mobile_number is not None else contact.phone_number
            contact.whatsapp_no = alternate_mobile_number if alternate_mobile_number is not None else contact.whatsapp_no
            contact.contact_type = contact_type_
            # Save the updated contact
            contact.save()
            data = contact
            success = True
        except ValidationError as e:
            errors = [str(e)]
        except Exception as e:
            print(e)
            errors = [str(e)]
    else:
        try:
            # Create a new contact if no id is provided
            contact = ContactDetalis.objects.create(
                contact_person_name=name,
                email=email,
                phone_number=mobile_number,
                whatsapp_no=alternate_mobile_number,
                contact_type=contact_type_
            )
            data = contact
            success = True
        except ValidationError as e:

            errors = [str(detail) for detail in e.detail]
        except Exception as e:
             
            errors = [str(e)]
    return {"data": data, "success": success, "errors": errors}


def CreateAccountGeneralLedger(data, isValidations):
    """
        Depence on condition created the GeneralLedger
    """
    errors = []
    required_field = ['date', "DebitAccount", "amount", "created_by",'creditAccount']
    voucher_no_fields = ['payment_voucher_no']
    type_for_voucher_no = {"payment_voucher_no": "payment_voucher"}
    voucher_type = None

    for field in required_field:
        if not str(data.get(field, '')).strip() or data.get(field, '') == None:
            errors.append(f'{field} is required')

    # Voucher No Logic
    filled_voucher_fields = [field for field in voucher_no_fields if data.get(field)]
    if not filled_voucher_fields:
        errors.append("Voucher No Field is required")
    elif len(filled_voucher_fields) > 1:
        errors.append("More than one Voucher No is provided. Only one is allowed.")
    else:
        voucher_type = type_for_voucher_no.get(filled_voucher_fields[0])
        if not voucher_type:
            errors.append("Voucher Type is not available.")

    # Either customer_supplier or employee
    if data.get('customer_supplier') and data.get('employee'):
        errors.append('Only one of Customer/Supplier or Employee should be filled')
    if not data.get('customer_supplier') and not data.get('employee'):
        errors.append('Either Customer/Supplier or Employee must be filled')
    
    # Account logic
    try:
        DebitAccount = AccountsMaster.objects.get(id=data['DebitAccount'])
    except AccountsMaster.DoesNotExist:
        errors.append("Invalid Debit Account selected.")
        return {"success": False, "errors": errors, "data": ""}
    try:
        CreditAccount = AccountsMaster.objects.get(id=data['creditAccount'])
    except AccountsMaster.DoesNotExist:
        errors.append("Invalid Debit Account selected.")
        return {"success": False, "errors": errors, "data": ""}
    
    try:
        if data['amount'] == None:
            data['amount'] = 0
        amount = Decimal(str(data['amount']).strip())
    except (ValueError, TypeError):
        errors.append("Amount must be a valid number.")
        return {"success": False, "errors": errors, "data": ""}

    """Update the debit"""
    if DebitAccount.account_type == "Employee":
        if DebitAccount.accounts_group_name and str(DebitAccount.accounts_group_name.accounts_group_name).lower() == "current assets":
            data['debit'] = amount


    """Update the credit"""
    if CreditAccount.account_type in ["Bank", "Cash"]:
        data['credit'] = amount
    if (data.get('debit') != None and data['debit'] <= 0):
        errors.append('Debit  must be filled')
    if (data.get('credit') != None and data['credit'] <= 0):
        errors.append('Credit must be filled')

    if errors:
        error = "".join(errors) 
        return {"success": False, "errors": error, "data": ""}

    if isValidations:
        return {"success": True, "errors": [], "data": ""}

    # Create ledger
    if data.get('debit'):
        ledger = AccountsGeneralLedger.objects.create(
            date=data['date'],
            account=DebitAccount,
            debit=data.get('debit'),
            customer_supplier_id=data.get('customer_supplier'),
            employee_id=data.get('employee'),
            remark=data.get('remark', ''),
            created_by_id=data['created_by'],
            payment_voucher_no_id=data.get('payment_voucher_no'),
            voucher_type=voucher_type,
        )
    if data.get('credit'):
         ledger = AccountsGeneralLedger.objects.create(
            date=data['date'],
            account=CreditAccount,
            credit=data.get('credit'),
            customer_supplier_id=data.get('customer_supplier'),
            employee_id=data.get('employee'),
            remark=data.get('remark', ''),
            created_by_id=data['created_by'],
            payment_voucher_no_id=data.get('payment_voucher_no'),
            voucher_type=voucher_type,
        )

    return {"success": True, "errors": [], "data": ledger}


def ValidationForeignKeys(app, model, id):
    """
        Comman Validate the ForeignKeys
    """
    error = []
    try:
        Model_Class = apps.get_model(app_label=app, model_name=model)
        instance = Model_Class.objects.get(id=id) 
        return {"success": True, "error": [], "instance": instance}
    except ObjectDoesNotExist:
        return {"success": False, "error": f"{model} with id {id} does not exist."}
    except LookupError:
        return {"success": False, "error": f"Model '{model}' not found in app '{app}'."}
    except Exception as e:
        return {"success": False, "error": f"{model} raised an error: {e}"}


def getTaxDetails(itemmaster_id, states):
    item_master = ItemMaster.objects.get(id=itemmaster_id)
    states = str(states).upper()

    if item_master:
        hsn = item_master.item_hsn
        if states == "TAMIL NADU":
            sgst = int(hsn.gst_rates.rate) / 2
            cgst = int(hsn.gst_rates.rate) / 2
            return {"sgst": sgst, "cgst": cgst, "igst": ""}
        else:
            igst = int(hsn.gst_rates.rate)
            return {"sgst": "", "cgst": "", "igst": igst}


class CommonError(Exception):
    pass

def cal_total_dic_value(data):
    print(data)
    if data:
        return (sum(Decimal(v) for v in data.values()) or 0)
    else :
        return Decimal(0)

def CalculateInventoryQty(queryset):
    update_list = []
    for item in queryset:
        total_qty = sum(inventory.qty for inventory in item.inventory_id.all())
        item.qtyofInventoryApproval = total_qty if total_qty else None  # Update the attribute directly
        update_list.append(item)
    return update_list


def validate_with_serializer(data, model, serializer_class, model_related_key,
                             model_related_field, name_field, related_field_name):
    errors = []
    for itemDetail in data:  # Iterate through the item details or other income charges

        # Ensure the related field exists in itemDetail

        if model_related_field not in itemDetail:
            errors.append(f"Missing related field {model_related_field} in {name_field}.")
            continue
        try:
            # Get the related item (e.g., ItemMaster or OtherIncomeCharges)
            related_item = model_related_key.objects.get(id=itemDetail[model_related_field])

        except model_related_key.DoesNotExist:
            errors.append(f"{name_field} with ID {itemDetail[model_related_field]} not found.")
            continue

        # Check if there's an ID in the item detail
        if 'id' in itemDetail and itemDetail['id']:  # Check if 'id' exists in itemDetail

            item_instance = model.objects.filter(id=itemDetail['id']).first()

            if not item_instance:
                errors.append(f"{name_field}   not found.")
                continue  # Proceed to next item if an error is found
            else:
                serializer = serializer_class(item_instance, data=itemDetail, partial=True)
        else:
            serializer = serializer_class(data=itemDetail)

        # Validate the serializer and capture errors
        if not serializer.is_valid():
            try:
                for field, error_list in serializer.errors.items():
                    errors.append(
                        f"{field}: {'; '.join([str(e) for e in error_list])}"
                    )


            except Exception as e:
                print("e", e)
                errors.append(str(e))
                # errors.append(str(e)) 
    return errors


def validate_with_serializer_with_out_related_model(data, model, serializer_class, name_field, related_field_name):
    errors = []

    for itemDetail in data:  # Iterate through the item details or other income charges

        # Ensure the related field exists in itemDetail

        # Check if there's an ID in the item detail
        if 'id' in itemDetail and itemDetail['id']:  # Check if 'id' exists in itemDetail

            item_instance = model.objects.filter(id=itemDetail['id']).first()
            if not item_instance:
                errors.append(f"{itemDetail[name_field]}   not found.")
                continue  # Proceed to next item if an error is found
            else:
                serializer = serializer_class(item_instance, data=itemDetail, partial=True)
        else:
            serializer = serializer_class(data=itemDetail)

        # Validate the serializer and capture errors
        if not serializer.is_valid():
            try:
                for field, error_list in serializer.errors.items():
                    errors.append(
                        f" {itemDetail[name_field]}-{field}: {'; '.join([str(e) for e in error_list])}"
                    )
            except Exception as e:
                errors.append(str(e))
    return errors


def deleteCommanLinkedTable(prev_ids, new_ids, model):
    errors = []
    ids = [id for id in prev_ids if id not in new_ids]
    print("ids", ids)

    if len(ids) > 0:
        try:
            # Retrieve the model instances using filter
            model_instances = model.objects.filter(id__in=ids)
            if model_instances.exists():
                # Delete each model instance
                for modeldata in model_instances:
                    modeldata.delete()

            else:
                errors.append('No records found to delete.')

        except ProtectedError:

            errors.append('This data is linked with other modules and cannot be deleted.')
        except Exception as e:
            errors.append(str(e))
        return errors
    print("---==")
    return []


def commanSaveAllModel(model, kwargs, serializer, model_name):
    errors = []
    instance_ = None
    success = False

    if kwargs.get('id'):
        try:
            model_instance = model.objects.get(id=kwargs['id'])
            serializer_ = serializer(model_instance, data=kwargs, partial=True)
        except model.DoesNotExist:
            errors.append(f"{model_name} with id {kwargs['id']} not found.")
            return {"success": success, "errors": errors, "instance": instance_}
    else:
        serializer_ = serializer(data=kwargs)
    if serializer_.is_valid():
        try:
            instance_ = serializer_.save()
            success = True
        except Exception as e:
            errors.append(str(e))
    else:
        errors = [f"{field}: {'; '.join([str(e) for e in error])}" for field, error in serializer_.errors.items()]

    return {"success": success, "errors": errors, "instance": instance_}


def SendMail(access_token, to_emails, subject, body_content, cc=[], bcc=[], content_type_data="HTML",attachments=None):
    try:
        success = False
        error = []
        # Microsoft Graph API endpoint for sending emails
        graph_api_url = 'https://graph.microsoft.com/v1.0/me/sendMail'
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/json'
        }
        # Ensure `to_emails`, `cc`, and `bcc` are lists of plain strings, not stringified lists
        to_emails = json.loads(to_emails) if isinstance(to_emails, str) else to_emails
        cc = json.loads(cc) if isinstance(cc, str) else cc
        bcc = json.loads(bcc) if isinstance(bcc, str) else bcc

        # Build the email message
        email_data = {
            "message": {
                "subject": subject,
                "body": {"contentType": content_type_data, "content": body_content},
                "toRecipients": [{"emailAddress": {"address": email.strip()}} for email in to_emails if email],
                "ccRecipients": [{"emailAddress": {"address": email.strip()}} for email in cc if email],
                "bccRecipients": [{"emailAddress": {"address": email.strip()}} for email in bcc if email],
                "attachments": attachments
            }
        }

        # Send the email via Microsoft Graph API
        response = requests.post(graph_api_url, headers=headers, json=email_data)

        # Check the response
        if response.status_code == 202:  # 202 means accepted
            success = True
            error = []
        else:
            error.append(f"Error sending email: {response.status_code} - {response.text}")

        return {"success": success, "error": error}

    except Exception as e:
        return {"success": False, "error": [str(e)]}


from functools import wraps
from graphql import GraphQLError
from rest_framework.response import Response
from rest_framework import status

def permission_required(models, type="query", action="view"):
    
    """ it is only for view permission """
    def decorator(func):
        @wraps(func)
        def wrapper(root, info, *args, **kwargs):
            user = info.context.user
            if not user or not user.is_authenticated:
                raise GraphQLError(
                    "Authentication credentials were not provided.",
                    extensions={"code": "UNAUTHENTICATED"})

            if type == "query":
                has_any_permission = any(has_permission_query(user, model, action) for model in models)
                if not has_any_permission:
                    raise GraphQLError(
                        f"You do not have permission to view any of these models: {', '.join(models)}",
                        extensions={
                            "code": "PERMISSION_DENIED",
                            "missing_models": models
                        }
                    )

            elif type == "mutation":
                missing_permissions = []
                for model in models:
                    if not has_permission_mutations(user, model, action):
                        missing_permissions.append(model)

                if missing_permissions:
                    raise GraphQLError(
                        f"You do not have permission to {action} the following models: {', '.join(missing_permissions)}",
                        extensions={
                            "code": "PERMISSION_DENIED",
                            "missing_models": missing_permissions
                        }
                    )
            else:
                raise GraphQLError(
                    "Invalid permission type specified.",
                    extensions={"code": "INVALID_PERMISSION_TYPE"}
                )

            return func(root, info, *args, **kwargs)
        return wrapper
    return decorator

def mutation_permission(model_name, create_action, edit_action, edit_extra_actions=[]):
    """Check permission for create and edit"""
    def decorator(func):
        @wraps(func)
        def wrapper(self, info, *args, **kwargs):
            user = info.context.user
            instance_id = kwargs.get("id")
            action = None
            if instance_id:
                if len(edit_extra_actions) > 0:
                    action = edit_extra_actions
                else:
                    action = [edit_action]
            else:
                action = [create_action]
            if not user or not user.is_authenticated:
                raise GraphQLError(
                    "Authentication required",
                    extensions={"code": "UNAUTHENTICATED"}
                )
            if not has_permission_mutations(user, model_name, action):
                raise GraphQLError(
                    f"You do not have permission to {action} {model_name}",
                    extensions={"code": "PERMISSION_DENIED"}
                )

            return func(self, info, *args, **kwargs)
        return wrapper
    return decorator

def status_mutation_permission(model_name):
    """Check permission depence on status"""
    def decorator(func):
        @wraps(func)
        def wrapper(self, info, *args, **kwargs):
            user = info.context.user
            status = kwargs.get("status")
            if not user or not user.is_authenticated:
                raise GraphQLError(
                    "Authentication required",
                    extensions={"code": "UNAUTHENTICATED"}
                )
            if not has_permission_mutations(user, model_name, status):
                raise GraphQLError(
                    f"You do not have permission to {status} {model_name}",
                    extensions={"code": "PERMISSION_DENIED"}
                )
            return func(self, info, *args, **kwargs)
        return wrapper
    return decorator

def permission_required_mutations(models, action):
    def decorator(func):
        @wraps(func)
        def wrapper(root, info, *args, **kwargs):
            user = info.context.user

            if not user or not user.is_authenticated:
                raise GraphQLError("Authentication credentials were not provided.")

            for model in models:
                if not has_permission_query(user, model):
                    raise GraphQLError(f"You do not have permission to view {model}")

            return func(root, info, *args, **kwargs)
        return wrapper
    return decorator


def has_permission_mutations(user, model_name, action):
    """checking permission in mutations"""
    if not user or not user.is_authenticated:
        return False 
    try:
        user_management = UserManagement.objects.filter(user=user).first()
        if user_management and user_management.profile:
            allowed_perms = user_management.profile.allowed_permission.filter(model_name=model_name)
            for perm in allowed_perms:
                # Handle single string or list of actions
                if isinstance(action, str):
                    if perm.permission_options.filter(options_name=action).exists():
                        return True
                elif isinstance(action, (list, tuple, set)):
                    if perm.permission_options.filter(options_name__in=action).exists():
                        return True
        return False
    except Exception as e:
        print("Permission check error:", e)
        return False


def has_permission_query(user, model_list, action):
    if not user or not user.is_authenticated:
        return False
    try:
        user_management = UserManagement.objects.filter(user=user).first()
        if user_management and user_management.profile:
            if isinstance(model_list, str):
                model_list = model_list.split("__")
            allowed_perms = user_management.profile.allowed_permission.filter(
                model_name__in=model_list
            )
            for perm in allowed_perms:
                if perm.permission_options.filter(options_name__iexact=action).exists():
                    return True
        return False
    except Exception as e:
        print("Permission check error (query):", e)
        return False

def api_permission_required(models, type="query", action="view"):
    def decorator(func):
        @wraps(func)
        def wrapper(self, request, *args, **kwargs):
            user = request.user
            if not user or not user.is_authenticated:
                return Response({
                    "error": "Authentication credentials were not provided."
                }, status=status.HTTP_401_UNAUTHORIZED)

            has_any_permission = any(has_permission_query(user, model, action) for model in models)
            if not has_any_permission:
                return Response({
                    "error": f"You do not have permission to {action} any of these models: {', '.join(models)}",
                    "missing_models": models
                })

            return func(self, request, *args, **kwargs)
        return wrapper
    return decorator

def stockHistoryUpdate(actions, store, partCode, previous, updated, added, reduced, savedUserInstance, transActionModel,
                       transActionId, displayId, displayName, stock_link, conference_id=None):
     
    try:
        history_instance = StockHistory.objects.create(action=actions, store_link=store,
                            part_number=partCode, previous_state=str(previous),
                            updated_state=str(updated), added=str(added),
                            reduced=reduced, saved_by=savedUserInstance,
                            transaction_module=transActionModel,
                            transaction_id=transActionId, display_id=displayId,
                            display_name=displayName, stock_link=stock_link,
                            conference=conference_id)
        history_instance.save() 
    except Exception as e:
        print("error in save", e)


def CheckStockData(itemmaster , store, batchs, serial, required_stock):
    success = False
    errors = []
    needStock = []
    if itemmaster.batch_number:
        for batch in batchs:
            if batch.is_stock_reduce:
                continue
            batch_instance = batch.batch
            stock_data_existing = ItemStock.objects.filter(
                            part_number=itemmaster.id,
                            store=store.id,
                            batch_number=batch_instance.id
                        ).first()
            batch_lable = batch_instance.batch_number_name 
            if stock_data_existing:
                current_stock = Decimal(stock_data_existing.current_stock)
                required_stock= (batch.qty or 0)
            
                if (current_stock - required_stock) < 0:
                    needStock.append({
                                    "partcode": itemmaster.item_part_code,
                                    "batch": batch_lable,
                                    "needStock": abs(current_stock - required_stock)
                                })
            else:
                needStock.append({
                    "partcode": itemmaster.item_part_code,
                    "batch": batch,
                    "needStock": required_stock
                })
                
    elif itemmaster.serial:
        serial_ids = [serial.id for serial in serial]
       
        stock_data_existing = ItemStock.objects.filter(
                        part_number=itemmaster.id,
                        store=store.id,
                    ).first()
        if stock_data_existing:
            
            # Fetch the serial numbers associated with this stock
            serial_list = [serial.id for serial in stock_data_existing.serial_number.all()]

            # Check if all serial_ids are in serial_list
            missing_serial_ids = [serial_id for serial_id in serial_ids if serial_id not in serial_list]

            # Fetch the missing serial numbers
            missing_serials = SerialNumbers.objects.filter(id__in=missing_serial_ids) 
            if missing_serial_ids:
                # Extract serial_number for missing serial IDs
                needSerial = [serial.serial_number for serial in missing_serials]
                needStock.append({
                    "partcode": itemmaster.item_part_code,
                    "Serial": needSerial,
                    "needStock": len(needSerial)
                })
    else:
        stock_data_existing = ItemStock.objects.filter(part_number=itemmaster.id,
                                            store=store.id).first()
         
        if stock_data_existing:
            current_stock = Decimal(stock_data_existing.current_stock)
            required_stock = Decimal(required_stock)
            if (current_stock - required_stock) < 0:
                needStock.append({
                    "partcode": itemmaster.item_part_code,
                    "needStock": abs(current_stock - required_stock)
                })
        else:
            needStock.append({
                "partcode": itemmaster.item_part_code,
                "needStock": required_stock
            })

    if not needStock:
        success = True
    return {"success":success, "errors":errors, "needStock":needStock}


""" pdf created functions"""
#Add Empty Rows
def insert_empty_row_after(table, row_idx):
    # Get the reference row to copy the structure from
    ref_row = table.rows[row_idx]
    tr = ref_row._tr  # Access the underlying XML of the row
    new_tr = OxmlElement('w:tr')  # Create a new row element

    # Copy each cell from the reference row to the new row
    for cell in ref_row.cells:
        new_tc = OxmlElement('w:tc')  # Create a new cell
        new_tcPr = cell._tc.get_or_add_tcPr()  # Copy cell properties (e.g., width, borders)
        new_tc.append(new_tcPr)
        new_tc.append(OxmlElement('w:p'))  # Add an empty paragraph
        new_tr.append(new_tc)

    tr.addnext(new_tr)  # Insert the new row after the reference row
    return table.rows[row_idx + 1]  # Return the inserted row object

#Set Style Inside the Table Cells,if we already have the style that style will be applied else Aptos will be applied
def set_cell_font_style(table):
    """
    Applies default font name and size to table cell runs,
    but only if those properties were not already set in the Word template.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    # Only apply if size is not set (to preserve template formatting)
                    if run.font.size is None:
                        run.font.size = Pt(8)
                    
                    # Only apply if font name is not set (to preserve template formatting)
                    if run.font.name is None:
                        run.font.name = 'Aptos'

                    # ✅ Do NOT touch color, bold, italic, underline, or highlight

#Styling Outside the Tables like item Quotation No ,address etc ,if we already have the style that style will be applied else Aptos will be applied
def replace_text_in_run(run, key, value):
    # This Function is used to replace the place holder which is outside the table like paragraph
    placeholder = f"{{{{{key}}}}}"
    # print("placeholder",placeholder,run.text)
    if placeholder in run.text:
        run.text = run.text.replace(placeholder, value)  # Replace placeholder with actual value
        # print("placeholder, value",placeholder, value)
        if not run.font.size:
            run.font.size = Pt(8)
        if not run.font.name:
            run.font.name = 'Aptos'

#Styling the placeholder like Quotation No address Etc,if we already have the style that style will be applied else Aptos will be applied
def replace_placeholder_preserving_style(paragraph, key, value):
    """
    Replace {{key}} in a paragraph while preserving the formatting (font, bold, color, etc.)
    even if the placeholder is split across multiple runs.
    """
    placeholder = f"{{{{{key}}}}}"  
    full_text = ''.join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return

    # Combine all text into one string and store formatting
    new_text = full_text.replace(placeholder, value)

    # Use the style of the first run for the replacement
    first_run = paragraph.runs[0] if paragraph.runs else None

    # Clear all runs (only text, not paragraph object)
    for run in paragraph.runs:
        run.text = ""

    # Add the replacement text using preserved style
    new_run = paragraph.add_run(new_text)
    if first_run:
        new_run.font.name = first_run.font.name
        new_run.font.size = first_run.font.size
        new_run.bold = first_run.bold
        new_run.italic = first_run.italic
        new_run.underline = first_run.underline
        new_run.font.color.rgb = getattr(first_run.font.color, 'rgb', None)
        if getattr(first_run.font.highlight_color, "value", None):
            new_run.font.highlight_color = first_run.font.highlight_color

#Based on the aligment mentioned in mock data for the particular style will be applied here 
def apply_cell_style(cell, value, alignment="center"):
    """
    Safely sets value and alignment without removing existing styles from template.
    """
    paragraph = cell.paragraphs[0]

    # Use existing run if present to retain formatting (bold, italic, etc.)
    if paragraph.runs:
        run = paragraph.runs[0]
        run.text = str(value)  # Replace only text
    else:
        run = paragraph.add_run(str(value))  # Add new run if empty
        if not run.font.size:
            run.font.size = Pt(8)
        if not run.font.name:
            run.font.name = 'Aptos'

    # Set alignment (text alignment is paragraph-level, not run-level)
    alignment = alignment.lower()
    if alignment == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def fill_document_with_mock_data(doc: Document, mock_data: dict):
    table_names = mock_data.get("Table Name", [])# List of table placeholders to replace (like Item Table)
    other_tables = mock_data.get("Other Table", []) # The Table Which as no Header but want to add as a row should be in the Other Table

    for para in doc.paragraphs:
        for key, val in mock_data.items():
            if isinstance(val, str):
                for run in para.runs:
                    replace_text_in_run(run, key, val)

    for table in doc.tables:
        # doc.add_paragraph("      ")  # creates spacing

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if cell.text: 
                    for key, val in mock_data.items():
                        # If value is a string, replace any text placeholders in table cells
                        if isinstance(val, str) and f"{{{{{key}}}}}" in cell.text:
                            for paragraph in cell.paragraphs:
                                replace_placeholder_preserving_style(paragraph, key, val)
                        
                        # If value is a list (i.e., table data), check for a matching table name
                        elif isinstance(val, list):
                            for table_name in table_names:
                                placeholder = f"{{{{{table_name}}}}}"
                                if placeholder in cell.text and key == f"{table_name}_Datas":
                                    cell.text = "" # Clear the placeholder cell
                                    column_mapping = mock_data.get(f"{table_name}_Columns", {})
                                    data_list = val
                                    # If table is in 'Other Table' list, insert raw key-value data which has no headers in the template
                                    if table_name in other_tables:
                                        if not data_list:
                                            table._tbl.remove(row._tr)
                                            continue

                                        headers = list(data_list[0].keys())
                                        insert_row_idx = row_idx  # Start with current row index

                                        # First item goes in the current row
                                        # if data_list:
                                        #     for col_idx, header in enumerate(headers):
                                        #         if col_idx < len(row.cells):
                                        #             row.cells[col_idx].text = str(data_list[0].get(header, ""))
                                        
                                        style_key = f"{table_name}_Style"
                                        alignment_map = mock_data.get(style_key, {})
                                        if data_list:
                                            for col_idx, header in enumerate(headers):
                                                if col_idx < len(row.cells):
                                                    alignment = alignment_map.get(header, "center")
                                                    value = data_list[0].get(header, "")
                                                    apply_cell_style(row.cells[col_idx], value, alignment)

                                        # Additional items get new rows
                                        for item in data_list[1:]:
                                            new_row = insert_empty_row_after(table, insert_row_idx)
                                            for col_idx, header in enumerate(headers):
                                                if col_idx < len(new_row.cells):
                                                    # new_row.cells[col_idx].text = str(item.get(header, ""))
                                                    alignment = alignment_map.get(header, "center")
                                                    value = item.get(header, "")
                                                    apply_cell_style(new_row.cells[col_idx], value, alignment)
                                            insert_row_idx += 1
                                        set_cell_font_style(table)
                                    # Handle mapped or nested header tables
                                    else:
                                        headers = []
                                        # Check for nested header format like: CGST[%], SGST[Amount], etc.
                                        is_nested = any("[" in key and "]" in key for key in column_mapping.keys())
                                        # Case: Nested header from 2 rows above
                                        if is_nested and row_idx - 2 >= 0:
                                            parent_row = table.rows[row_idx - 2]
                                            child_row = table.rows[row_idx - 1]
                                            #Example Nested Header look like ['S.No[S.No]', 'HSN/SAC[HSN/SAC]', 'Taxable Value[Taxable Value]', 'CGST[%]', 'CGST[Amount]', 'SGST[%]', 'SGST[Amount]', 'IGST[%]', 'IGST[Amount]'] Should mention in the Columns
                                            for i, cell in enumerate(child_row.cells):
                                                parent_text = parent_row.cells[i].text.strip() if i < len(parent_row.cells) else ""
                                                child_text = cell.text.strip()
                                                if parent_text:
                                                    full_header = f"{parent_text}[{child_text}]"
                                                else:
                                                    full_header = child_text
                                                headers.append(full_header)

                                        # Fallback: Single header row just above
                                        elif row_idx - 1 >= 0:
                                            # Normal Header Look  Like ['S.No', 'Description', 'HSN', 'Qty', 'UOM', 'Rate', 'Disc', '%', 'Amount', 'Total']
                                            header_row = table.rows[row_idx - 1]
                                            headers = [c.text.strip() for c in header_row.cells]

                                        # Fill table with data based on resolved headers
                                        if data_list and headers:
                                            first_item = data_list[0]
                                            for col_idx, header in enumerate(headers):
                                                data_key = column_mapping.get(header)
                                                if data_key:
                                                    # row.cells[col_idx].text = str(first_item.get(data_key, ""))
                                                    style_key = f"{table_name}_Columns_Style"
                                                    alignment_map = mock_data.get(style_key, {})
                                                    alignment = alignment_map.get(header, "center")
                                                    value = first_item.get(data_key, "")
                                                    apply_cell_style(row.cells[col_idx], value, alignment)
                                            insert_row_idx = row_idx
                                            for item in data_list[1:]:
                                                insert_row_idx += 1
                                                new_row = insert_empty_row_after(table, insert_row_idx - 1)
                                                for col_idx, header in enumerate(headers):
                                                    data_key = column_mapping.get(header)
                                                    if data_key:
                                                        # new_row.cells[col_idx].text = str(item.get(data_key, ""))
                                                        style_key = f"{table_name}_Columns_Style"
                                                        alignment_map = mock_data.get(style_key, {})
                                                        alignment = alignment_map.get(header, "center")
                                                        value = item.get(data_key, "")
                                                        apply_cell_style(new_row.cells[col_idx], value, alignment)
                                            set_cell_font_style(table)
                                        break 
    return doc


def format_indian_currency(amount, symbol='₹', remove_symbol=False):
    amount = float(amount)
    parts = f"{amount:.2f}".split(".")
    integer_part = parts[0]
    decimal_part = parts[1]

    # First 3 digits (right to left), then pairs of 2 digits
    if len(integer_part) > 3:
        last_three = integer_part[-3:]
        remaining = integer_part[:-3]
        remaining = ",".join([remaining[max(i - 2, 0):i] for i in range(len(remaining), 0, -2)][::-1])
        formatted = remaining + "," + last_three
    else:
        formatted = integer_part

   
    if remove_symbol:
       
        return f"{formatted}.{decimal_part}"
    else:
        return f"{symbol}{formatted}.{decimal_part}" if symbol else f"{formatted}.{decimal_part}"


def format_international_currency(amount,symbol='$',remove_symbol = False): 
    if remove_symbol:
        return f"{float(amount):,.2f}"
    else:
        return f"{symbol}{float(amount):,.2f}" if symbol else f"{float(amount):,.2f}"

def format_currency(amount, symbol, remove_symbol=False):
    amount_ = 0  if amount == "-" else amount 
    if symbol == '₹':
        return format_indian_currency(amount_,symbol,remove_symbol )
    else:
        return format_international_currency(amount_, symbol, remove_symbol)


def createPdfHsnTableContent(itemdetails, othercharges, currency_=None):
        currency = currency_
        aggregated_data = {}
        def safe_float(value):
            return float(value) if value not in [None, "null"] else 0.0

            # Aggregate item details
        for item in itemdetails.all():
            hsn = item.hsn.hsn_code
            rate = safe_float(item.rate)
            qty = safe_float(item.qty)
            disc_val = safe_float(item.discount_value)
            disc_pct = safe_float(item.discount_percentage)

            discount_amt = (disc_pct / 100) * (rate * qty) if disc_pct > 0 else disc_val
            taxable_value = (rate * qty) - discount_amt

            cgst = safe_float(item.cgst)
            sgst = safe_float(item.sgst)
            igst = safe_float(item.igst)

            if hsn not in aggregated_data:
                aggregated_data[hsn] = {
                    "taxable": 0.0,
                    "cgst": cgst,
                    "sgst": sgst,
                    "igst": igst
                }

            aggregated_data[hsn]["taxable"] += taxable_value

            # Aggregate other income charges
        
        if   othercharges.exists():
            for charge in othercharges.all():
                hsn = charge.other_income_charges_id.hsn.hsn_code
                amount = safe_float(charge.amount)
                cgst = safe_float(charge.cgst)
                sgst = safe_float(charge.sgst)
                igst = safe_float(charge.igst)

                if hsn not in aggregated_data:
                    aggregated_data[hsn] = {
                        "taxable": 0.0,
                        "cgst": cgst,
                        "sgst": sgst,
                        "igst": igst
                    }

                aggregated_data[hsn]["taxable"] += amount

        # Final tax list
        list_tax = []

        def remove_number(data):
            value = float(data)
            return int(value) if value.is_integer() else value



        for i, (hsn, data) in enumerate(aggregated_data.items(), start=1):
            taxable = data["taxable"]
            cgst = data["cgst"]
            sgst = data["sgst"]
            igst = data["igst"] 
            list_tax.append({
                        "S.No": str(i),
                        "HSN/SAC": str(hsn),
                        "Taxable Value": format_currency(f"{taxable:.2f}",currency ,False),
                        "cgst %": f"{remove_number(cgst)}" if cgst else "-",
                        "cgst Amount": (format_currency(f"{(cgst * taxable / 100):.2f}",currency,False) if cgst else "-"),
                        "sgst %": f"{remove_number(sgst)}" if sgst else "-",
                        "sgst Amount":(format_currency(f"{(sgst * taxable / 100):.2f}",currency ,False) if sgst else "-"),
                        "igst %": f"{remove_number(igst)}" if igst else "-",
                        "igst Amount": (format_currency(f"{(igst * taxable / 100):.2f}",currency ,False) if igst else "-")
                    })
        return list_tax



def SaveTdsTcsEffective_date(tds=None, tcs=None,Effective_date= None,
                    individual_with_pan=None,other_with_pan=None,
                      other_without_pan=None,created_by=None):
    user = None
    if created_by:
        user = User.objects.filter(id=created_by).first()
    try:
        if tds and Effective_date:
            TdsTcs = TdsTcsEffectiveDate.objects.create(tds=tds, percent_individual_with_pan=individual_with_pan,
                                            percent_other_with_pan=other_with_pan, effective_date=Effective_date, created_by=user)
            TdsTcs.save()
        elif tcs and Effective_date:
            TdsTcs = TdsTcsEffectiveDate.objects.create(tcs=tcs, percent_individual_with_pan=individual_with_pan,
                            percent_other_with_pan=other_with_pan,  
                            percent_other_without_pan = other_without_pan, effective_date=Effective_date, created_by=user)
            TdsTcs.save()
        return {"success":True, "errors":""}
    except Exception as e:
        return {"success":False, "errors":f"An exception occurred -{e}"}
        
    


def save_common_data(data_list, instance_list, serializer_class, combo_label_list, context=None):
    """
    Save a list of data using the given serializer. Supports both create and update.

    Args:
        data_list (list): List of data dictionaries.
        instance_list (list): List of model instances (used for updates).
        serializer_class (Serializer): The DRF serializer class.
        combo_label_list (list): Label used in error messages for each item (e.g., 'Product 1', 'Service A').
        context (dict, optional): Context passed to the serializer.

    Returns:
        dict: {
            "ids": List of saved instance IDs,
            "success": Boolean status,
            "error": List of error messages,
            "instance_list": List of saved instances
        }
    """
    saved_ids = []
    saved_instances = []
    errors = []

    try:
        for item_data, current_instance, item_label in zip(data_list, instance_list, combo_label_list):
            try:
                if item_data.get('id'):
                    if not current_instance:
                        errors.append(f"{item_label} - Instance is missing for update.")
                        break
                    serializer = serializer_class(
                        current_instance,
                        data=item_data,
                        partial=True,
                        context={'request': context}
                    )
                else:
                    serializer = serializer_class(
                        data=item_data,
                        context={'request': context}
                    )

                if serializer.is_valid():
                    serializer.save()
                    saved_ids.append(serializer.instance.id)
                    saved_instances.append(serializer.instance)
                else:
                    for field, field_errors in serializer.errors.items():
                        for err in field_errors:
                            errors.append(f"{item_label} - {field}: {err}")
                    break  # Stop on first validation error

            except ValidationError as e:
                errors.extend([f"{item_label} - {msg}" for msg in e.detail])
                break
            except Exception as e:
                errors.append(f"{item_label} - Unexpected save error: {str(e)}")
                break

    except Exception as e:
        errors.append(f"Unexpected error while processing {item_label}: {str(e)}")

    return {
        "ids": saved_ids,
        "success": len(errors) == 0,
        "error": errors,
        "instance_list": saved_instances,
    }


def validate_common_data(data_list, instance_list, serializer_class, combo_label_list, context=None):
    """
    Validate a list of data using the provided serializer.
    Supports both creation and update operations.

    Args:
        data_list (list): List of data dictionaries.
        instance_list (list): List of model instances (used for updates).
        serializer_class (Serializer): The DRF serializer class for validation.
        combo_label_list (list): Label used in error messages for each item (e.g., 'Service 1', 'Product A').
        context (dict, optional): Context passed to the serializer.

    Returns:
        list: A list of validation error messages, if any.
    """
    errors = [] 
    try:
        for item_data, current_instance, item_label in zip(data_list, instance_list, combo_label_list):
            
            try:
                if item_data.get('id', None):
                    if not current_instance:
                        errors.append(f"{item_label} - Instance is missing for validation.")
                        break
                    serializer = serializer_class(
                        current_instance,
                        data=item_data,
                        partial=True,
                        context={'request': context}
                    )
                else:
                    serializer = serializer_class(
                        data=item_data,
                        context={'request': context}
                    )

                if not serializer.is_valid():
                    for field, field_errors in serializer.errors.items():
                        for err in field_errors:
                            errors.append(f"{item_label} - {field}: {err}")
                    break  # Stop on first validation failure

            except Exception as e:
                errors.append(f"{item_label} - Unexpected validation error: {str(e)}")
                break

    except Exception as e:
        errors.append(f"Unexpected top-level validation error: {str(e)}")

    return errors


def validate_common_data_and_send_with_instance(data_list, instance_list, serializer_class, combo_label_list, context=None):
    """
    Validate a list of data using the provided serializer.
    Supports both creation and update operations.

    Args:
        data_list (list): List of data dictionaries.
        instance_list (list): List of model instances (used for updates).
        serializer_class (Serializer): The DRF serializer class for validation.
        combo_label_list (list): Label used in error messages for each item (e.g., 'Service 1', 'Product A').
        context (dict, optional): Context passed to the serializer.

    Returns:
        dict: {
            "error": list of error messages, if any;
            "instance": list of valid serializer instances.
        }
    """
    errors = []
    valid_instances = [] 
    try:
        for item_data, current_instance, item_label in zip(data_list, instance_list, combo_label_list):
             
            try:
                if item_data.get('id'):
                    if not current_instance:
                        errors.append(f"{item_label} - Instance is missing for validation.")
                        break

                    serializer = serializer_class(
                        current_instance,
                        data=item_data,
                        partial=True,
                        context={'request': context} if context else {}
                    )
                else:
                    serializer = serializer_class(
                        data=item_data,
                        context={'request': context} if context else {}
                    ) 
                if serializer.is_valid():
                    valid_instances.append(serializer)
                else:
                    for field, field_errors in serializer.errors.items():
                        for err in field_errors:
                            errors.append(f"{item_label} - {field}: {err}")
                    break  # Stop on first validation failure

            except Exception as inner_exc:
                errors.append(f"{item_label} - Unexpected validation error: {str(inner_exc)}")
                break

    except Exception as outer_exc:
        errors.append(f"Unexpected top-level validation error: {str(outer_exc)}")

    return {
        "error": errors,
        "instance": valid_instances
    }


"""rate allow more the maximum price but do not allow less than minimum price"""
def is_within_range(value, min_price, max_price):
    # return min_price <= value <= max_price
    return min_price <= value

def SellItemCheckRange(itemList, exchange_rate=1):
    out_of_range_list = []
    current_currency = None
    currenyItemmasterObjects = None
    errors = []
    for item in itemList:
        rate = item.after_discount_value_for_per_item if item.after_discount_value_for_per_item else item.rate
        itemmaster = None

        if type(item.itemmaster) != int:
            itemmaster = item.itemmaste
        else:
            itemmaster = ItemMaster.objects.get(id=item.itemmaster)

        part_code = itemmaster.item_part_code
        min_price = itemmaster.item_min_price
        max_price  = itemmaster.item_mrp
        if not min_price and not max_price:
            errors.append(f"{part_code} -> Min Price and Max Price Not Found")
            continue
        exchange_rate = Decimal(exchange_rate) if type(exchange_rate) == str else exchange_rate
        if exchange_rate > 1:
            min_price = itemmaster.item_min_price / exchange_rate
            max_price = itemmaster.item_mrp / exchange_rate
        else:
            min_price = itemmaster.item_min_price
            max_price = itemmaster.item_mrp
        
        if not (min_price <= rate <= max_price):
            errors.append(
                f"{part_code} per item rate {rate:.2f} is out of allowed range "
                f"({min_price:.2f} - {max_price:.2f})."
            )
            continue
        


    return {"success": len(errors) == 0, "error": errors}



class DecimalEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, Decimal):
            return float(obj)  # or str(obj) if you want string output
        return super().default(obj)

def company_info():
    cache_key = "company_Data"
    cached_data = cache.get(cache_key)
    
    if cached_data is not None:
        return cached_data
    company_master = CompanyMaster.objects.all().first() 
    cache.set(cache_key, company_master, timeout=86400)# cache for 14 hours
    return company_master

def clear_company_info_cache():
    cache_key = "company_Data"
    cached_data = cache.get(cache_key)
    if cached_data is not None:
        cache.delete("company_Data")
    
    cache_bank_key = "company_bank_detail"
    cached_bank_data = cache.get(cache_bank_key)
    if cached_bank_data is not None:
        cached_bank_data.delete("company_Data")

def company_bank_info():
    cache_key = "company_bank_detail"
    cached_data = cache.get(cache_key)
    
    if cached_data is not None:
        return cached_data
    company_bank_info = BankDetails.objects.filter(is_active=True).first()
    if not company_bank_info:
        return None
    cache.set(cache_key, company_bank_info, timeout=86400)# cache for 14 hours
    return company_bank_info

def get_conference_stock(id, store_id, conference_id):
    
    """ geting conference stock datas"""
    queryset = StockHistory.objects.filter(
        conference=conference_id,
        part_number__id=id,
        store_link__id=store_id
    ).annotate(
        added_float=Cast('added', FloatField()),
        reduced_float=Cast('reduced', FloatField())
    )
 
    totals = queryset.aggregate(
        total_added=Sum(
            Case(
                When(action='ADD', then='added_float'),
                default=Value(0.0),
                output_field=FloatField()
            )
        ),
        total_reduced=Sum(
            Case(
                When(action='DELETE', then='reduced_float'),
                default=Value(0.0),
                output_field=FloatField()
            )
        )
    )
    balance = (totals.get("total_added") or 0) - (totals.get("total_reduced") or 0)
    return balance
