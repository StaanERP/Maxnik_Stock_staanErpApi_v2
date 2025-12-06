import subprocess

# import pythoncom
# import win32com.client
from django.http import HttpResponse
from docx import Document
from datetime import datetime
from docx.shared import Pt
from num2words import num2words
from docx2pdf import convert
from decimal import Decimal, ROUND_DOWN
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
import json
import os
from Staanenquiryfromwithazure.settings import *
import platform
import base64
import json
import subprocess
import requests
import tempfile

def set_cell_font_style(table, Isquotation):
    for row in table.rows:
        for cell in row.cells:
            # Check if the cell contains the word "QUOTATION"
            if "QUOTATION" in cell.text:
                continue  # Skip font changes for this cell
            if "SalesOrder" in cell.text:
                continue  # Skip font changes for this cell
            # Set font size and style for other cells
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)  # Set font size to 8
                    run.font.name = 'Arial'  # Set font style to Arial


def generate_item_combo(row_cells, item_combo_details):
    display_name_dict = {}

    # Group item combos by display name
    for item in item_combo_details:
        display_name = item["display"]
        if display_name not in display_name_dict:
            display_name_dict[display_name] = []
        display_name_dict[display_name].append(item)

    # Iterate through grouped items and format text
    for display_name, items in display_name_dict.items():
        # Add display name with underline
        paragraph = row_cells[1].add_paragraph()
        run = paragraph.add_run(display_name)
        run.underline = True

        # Add item details below the display name
        for idx, item in enumerate(items, start=1):
            item_desc = f"\n{idx}. {item['itemmaster']['itemPartCode']} - {item['itemmaster']['itemName']} ({item['qty']} {item['uom']['name']})"
            paragraph.add_run(item_desc)  # Add a newline after each item


# function to generate item details
def generate_item_detail(table, items):
    try:
        item_details = items["items"][0]["itemDetails"]
        num_items = len(item_details)
        total_qty = 0
        total_discount = 0
        final_amount = 0
        for i in range(num_items):
            detail = item_details[i]
            item_master = detail["itemmaster"]
            overall_amount = float(detail["amount"])
            row_cells = table.add_row().cells

            # Set up each cell with appropriate text and alignment
            row_cells[0].text = str(i + 1)
            row_cells[1].text = detail["description"]

            if detail.get("itemComboItemDetails"):
                generate_item_combo(row_cells, detail["itemComboItemDetails"])

            row_cells[2].text = str(detail["hsn"]["hsnCode"])
            qty = int(float(detail["qty"]))
            total_qty += qty
            row_cells[3].text = str(qty)
            row_cells[4].text = detail["uom"]["name"]
            rate = int(float(detail['rate']))
            row_cells[5].text = f"{rate}"
            if detail['afterDiscountValueForPerItem']:
                discount = int(float(detail['rate']) - float(detail['afterDiscountValueForPerItem']))
                total_discount += discount
                row_cells[6].text = f"{discount}"
            else:
                row_cells[6].text = " "

            amount = float(detail["amount"])
            tax_percentage = detail['tax']
            tax_amount = (amount * tax_percentage) / 100
            row_cells[7].text = f"{tax_percentage}%"

            # Align tax amount right
            row_cells[8].text = f"{tax_amount:.2f}"

            # Align overall amount right
            row_cells[9].text = f"{overall_amount:.2f}"

            final_amount += overall_amount

            # Center align all cells except the description cell
            for idx, cell in enumerate(row_cells):
                if idx == 1:  # Skip the description cell
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif idx == 5 or idx == 8 or idx == 9:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add the summary row
        summary_row_cells = table.add_row().cells
        summary_row_cells[0].text = "Total"
        summary_row_cells[1].text = " "  # Empty description cell
        summary_row_cells[2].text = " "  # Empty HSN cell
        summary_row_cells[3].text = str(total_qty)  # Total quantity
        summary_row_cells[4].text = " "  # Empty UOM cell
        summary_row_cells[5].text = " "  # Empty rate cell
        summary_row_cells[6].text = str(total_discount)  # Total discount
        summary_row_cells[7].text = " "  # Empty tax percentage cell
        summary_row_cells[8].text = " "  # Empty tax amount cell
        # summary_row_cells[9].text = str(final_amount)

        #   Right align final amount
        summary_row_cells[9].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Set background color for the summary row
        for cell in summary_row_cells:
            cell._element.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            )

            # Center alignment for all cells in the summary row
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Make the final amount bold
        final_amount_cell = summary_row_cells[9]
        final_amount_paragraph = final_amount_cell.paragraphs[0]
        run = final_amount_paragraph.add_run(str(final_amount))  # Add the final amount as text
        run.bold = True  # Make it bold
        summary_row_cells[9].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        return final_amount
    except Exception as e:
        print(e)


# function to generate quotation details
def generate_quotation(table, datalist):
    # Extract data from datalist
    item = datalist["items"][0]
    customer_id = item["customerId"]["supplierNo"]
    customer_name = item["customerId"]["companyName"]
    customer_address = item["customerAddress"]
    contact_person = item["customerContactPerson"]
    quotation_no = item["quotationNo"]["linkedModelId"]
    quotation_date = datetime.fromisoformat(item["CreatedAt"].replace("Z", "+00:00")).strftime("%d-%m-%Y")
    sales_person = item["salesPerson"]["username"]
    department_name = item["department"]["name"]
    customer_gst = item["customerId"]["gstin"]

    # Construct the customer address string
    address_string = (f"{customer_address['addressLine1']} "
                      f"{customer_address['addressLine2']}\n"
                      f"{customer_address['city']}\n"
                      f"{customer_address['state']} - "
                      f"{customer_address['pincode']}")

    contact_person_string = (f"{contact_person['contactPersonName']}\n"
                             f"              {contact_person['phoneNumber']}\n"
                             f"              {contact_person['email']}\n")

    # Create a dictionary to map placeholders to values
    placeholders = {
        '{customerId}': customer_id,
        '{customerName}': customer_name,
        '{customerAddress}': address_string,
        '{contactPerson}': contact_person_string,
        '{quotationNo}': quotation_no,
        '{quotationDate}': quotation_date,
        '{salesPerson}': sales_person,
        '{departmentName}': department_name,
        '{customerGstIn}': customer_gst
    }

    # Replace placeholders in the first table
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            original_text = cell.text  # Store the original text
            for placeholder, value in placeholders.items():
                if placeholder in original_text:
                    updated_text = original_text.replace(placeholder, str(value))
                    paragraph = cell.paragraphs[0]
                    paragraph.clear()
                    if placeholder == '{customerName}' or placeholder == '{quotationNo}' or placeholder == '{customerId}':
                        run = paragraph.add_run(updated_text)
                        run.bold = True
                    else:
                        paragraph.add_run(updated_text)
                    break


def generate_salesorder(table, datalist):
    item = datalist["items"][0]
    buyer_id = item["buyer"]["supplierNo"]
    buyer_name = item["buyer"]["companyName"]
    buyer_address = item["buyerAddress"]
    contact_person = item["contact"]
    salesorder_no = item["salesOrderNo"]["linkedModelId"]
    salesorder_date = item["CreatedAt"]
    sales_person = item["salesPerson"]["username"]
    department_name = item["department"]["name"]
    buyer_gst = item["buyer"]["gstin"]

    address_string = (f"{buyer_address['addressLine1']} "
                      f"{buyer_address['addressLine2']}\n"
                      f"{buyer_address['city']}\n"
                      f"{buyer_address['state']} - "
                      f"{buyer_address['pincode']}")
    contact_person_string = (f"{contact_person['contactPersonName']}\n"
                             f"              {contact_person['phoneNumber']}\n"
                             f"              {contact_person['email']}\n")

    placeholders = {
        '{customerId}': buyer_id,
        '{customerName}': buyer_name,
        '{customerAddress}': address_string,
        '{contactPerson}': contact_person_string,
        '{salesOrderNo}': salesorder_no,
        '{salesOrderDate}': salesorder_date,
        '{salesPerson}': sales_person,
        '{departmentName}': department_name,
        '{buyerGstIn}': buyer_gst
    }

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            original_text = cell.text  # Store the original text
            for placeholder, value in placeholders.items():
                if placeholder in original_text:
                    updated_text = original_text.replace(placeholder, str(value))
                    paragraph = cell.paragraphs[0]
                    paragraph.clear()
                    if placeholder == '{customerName}' or placeholder == '{salesOrderNo}' or placeholder == '{customerId}':
                        run = paragraph.add_run(updated_text)
                        run.bold = True
                    else:
                        paragraph.add_run(updated_text)
                    break


def generate_amount(table, datalist, final_amount):
    # Extract net amount and other relevant totals from the JSON data
    net_amount = float(datalist["items"][0]["netAmount"])
    net_amount_in_words = num2words(net_amount, lang='en').replace("dollars",
                                                                   "rupees").title() + " Only"  # net_amount_in_words = num2words(net_amount, lang='en').replace("dollars", "rupees").title()

    item_total_before_tax = float(datalist["items"][0]["itemTotalBeforTax"])

    if datalist["items"][0]['other_charges_befor_tax'] != None:
        other_charges_befor_tax = datalist["items"][0]['other_charges_befor_tax']
        taxable_value = float(other_charges_befor_tax) + float(final_amount)
    else:
        taxable_value = float(final_amount)
    total_amount_after_tax = float(datalist["items"][0]["netAmount"])
    gst = float(datalist["items"][0]["taxTotal"])

    other_income_charges = datalist["items"][0]["otherIncomeCharge"]

    # Replace overall amounts in the document
    for row in table.rows:
        for cell in row.cells:

            # if "{totalAmoutInWords}" in cell.text:
            #     cell.text = cell.text.replace("{totalAmoutInWords}", net_amount_in_words)
            if "{totalAmount}" in cell.text:
                cell.text = cell.text.replace("{totalAmount}", f"{item_total_before_tax:.2f}")
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            if "{totalTax}" in cell.text:
                cell.text = cell.text.replace("{totalTax}", f"{taxable_value:.2f}")
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            if "{gst}" in cell.text:
                cell.text = cell.text.replace("{gst}", f"{gst:.2f}")
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            if "{AfterTax}" in cell.text:
                # Format the AfterTax value to 2 decimal places
                formatted_after_tax = f"{total_amount_after_tax:.2f}"

                # Replace the placeholder {AfterTax} with the formatted value
                cell.text = cell.text.replace("{AfterTax}", formatted_after_tax)

                # Find the newly inserted run and set it to bold
                run = cell.paragraphs[0].runs[-1]  # Get the last run in the paragraph
                run.bold = True

                # Align the entire paragraph to the right
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Process Other Income Charges
    current_income_index = 1  # Track the current index for other income charges
    for charge in other_income_charges:
        account_name = charge["otherIncomeChargesId"]["name"]
        percentage = charge["otherIncomeChargesId"]["hsn"]["gstRates"]["rate"]
        tax_rate = (float(charge["amount"]) * charge["otherIncomeChargesId"]["hsn"]["gstRates"]["rate"]) / 100
        amount = float(charge["amount"])

        # Replace placeholders in the table
        row_found = False  # Flag to check if a row was found and updated
        for row in table.rows:
            for cell in row.cells:
                # Check for placeholders and replace them
                if "{account}" in cell.text:
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    run = paragraph.add_run(account_name)  # Add the text as a run
                    run.bold = True  # Set the run to bold
                    row_found = True

                if "{%}" in cell.text:
                    cell.text = cell.text.replace("{%}", f"{percentage}%")
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    row_found = True

                if "{amount}" in cell.text:
                    cell.text = cell.text.replace("{amount}", str(tax_rate))
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    row_found = True

                if "{total}" in cell.text:
                    cell.text = cell.text.replace("{total}", f"{amount:.2f}")
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    row_found = True

            # If the row was found and updated, break to move to the next charge
            if row_found:
                current_income_index += 1  # Increment the index for the next entry
                break  # Exit the inner loop to avoid processing the same row multiple times
    # Check for placeholders and delete rows in a single loop
    for i in range(len(table.rows) - 1, -1, -1):  # Iterate in reverse
        row = table.rows[i]
        if any(
                "{account}" in cell.text or
                "{tax}" in cell.text or
                "{amount}" in cell.text
                for cell in row.cells
        ):
            row._element.getparent().remove(row._element)


def generate_amount_in_words(table, datalist):
    net_amount = float(datalist["items"][0]["netAmount"])
    net_amount_in_words = num2words(net_amount, lang='en').replace("dollars",
                                                                   "rupees").title() + " Only"
    for row in table.rows:
        for cell in row.cells:
            if "{totalAmoutInWords}" in cell.text:
                paragraph = cell.paragraphs[0]
                parts = cell.text.split("{totalAmoutInWords}")
                paragraph.clear()
                if parts[0]:
                    paragraph.add_run(parts[0])
                run = paragraph.add_run(net_amount_in_words)
                run.bold = True
                if len(parts) > 1 and parts[1]:
                    paragraph.add_run(parts[1])
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


# function to generate tax details like IGST CGST SGST
def generate_tax_data(table, items):
    item_details = items["items"][0]["itemDetails"]
    other_income_charges = items["items"][0]["otherIncomeCharge"]

    # Dictionary to store aggregated data
    aggregated_data = {}
    total_taxable_value = 0.0
    total_cgst_amount = 0.0
    total_sgst_amount = 0.0
    total_igst_amount = 0.0
    # Process item details
    for detail in item_details:
        hsn_code = detail["hsn"]["hsnCode"]
        rate = float(detail["rate"]) if detail["rate"] != "null" else 0.0
        qty = float(detail["qty"]) if detail["qty"] != "null" else 0.0

        # Handle discount values
        discount_value = float(detail["discountValue"]) if detail["discountValue"] != None else 0.0
        discount_percentage = float(detail["discountPercentage"]) if detail["discountPercentage"] != None else 0.0

        # Calculate the discount amount
        discount_amount = (discount_percentage / 100) * (rate * qty) if discount_percentage > 0 else discount_value

        # Calculate taxable value
        taxable_value = (rate - discount_amount) * qty

        # Initialize or update aggregated data for item details
        if hsn_code not in aggregated_data:
            aggregated_data[hsn_code] = {
                "total_taxable_value": taxable_value,
                "total_cgst_percentage": float(detail["cgst"]) if detail["cgst"] != None else 0.0,
                "total_sgst_percentage": float(detail["sgst"]) if detail["sgst"] != None else 0.0,
                "total_igst_percentage": float(detail["igst"]) if detail["igst"] != None else 0.0,
                "count": 1  # Count occurrences for future reference
            }
        else:
            # Sum the taxable value and percentages for item details
            aggregated_data[hsn_code]["total_taxable_value"] += taxable_value
            aggregated_data[hsn_code]["count"] += 1

    # Process other income charges
    if other_income_charges != []:
        for charge in other_income_charges:
            hsn_code = charge["otherIncomeChargesId"]["hsn"]["hsnCode"]
            amount = float(charge["amount"]) if charge["amount"] != "null" else 0.0

            # Extract GST rates from the nested structure
            gst_rates = charge["otherIncomeChargesId"]["hsn"]["gstRates"]
            cgst_rate = float(charge.get("cgst", 0)) if charge.get("cgst") != None else 0.0
            sgst_rate = float(charge.get("sgst", 0)) if charge.get("sgst") != None else 0.0
            igst_rate = float(charge.get("igst", 0)) if charge.get("igst") != None else 0.0

            # Initialize or update aggregated data for other income charges
            if hsn_code not in aggregated_data:
                aggregated_data[hsn_code] = {
                    "total_taxable_value": amount,
                    "total_cgst_percentage": cgst_rate,
                    "total_sgst_percentage": sgst_rate,
                    "total_igst_percentage": igst_rate,
                    "count": 1  # Count occurrences for future reference
                }
            else:
                # Sum the taxable value and percentages for other income charges
                aggregated_data[hsn_code]["total_taxable_value"] += amount

    # Now populate the table with aggregated data
    for i, (hsn_code, data) in enumerate(aggregated_data.items()):
        row_cells = table.add_row().cells

        # Serial Number (S NO)
        row_cells[0].text = str(i + 1)

        # HSN/SAC
        row_cells[1].text = str(hsn_code)

        # Taxable Value
        row_cells[2].text = f"{data['total_taxable_value']:.2f}"

        # CGST Amount Calculation
        cgst_amount = (data['total_cgst_percentage'] * data['total_taxable_value']) / 100
        if data['total_cgst_percentage'] > 0:
            row_cells[3].text = f"{data['total_cgst_percentage']:.2f}%"
            row_cells[4].text = f"{cgst_amount:.2f}"
        else:
            row_cells[3].text = "-"
            row_cells[4].text = "-"

        # SGST Amount Calculation
        sgst_amount = (data['total_sgst_percentage'] * data['total_taxable_value']) / 100
        if data['total_sgst_percentage'] > 0:
            row_cells[5].text = f"{data['total_sgst_percentage']:.2f}%"
            row_cells[6].text = f"{sgst_amount:.2f}"
        else:
            row_cells[5].text = "-"
            row_cells[6].text = "-"

        # IGST Amount Calculation
        igst_amount = (data['total_igst_percentage'] * data['total_taxable_value']) / 100
        if data['total_igst_percentage'] > 0:
            row_cells[7].text = f"{data['total_igst_percentage']:.2f}%"
            row_cells[8].text = f"{igst_amount:.2f}"
        else:
            row_cells[7].text = "-"
            row_cells[8].text = "-"

        total_taxable_value += data['total_taxable_value']
        total_cgst_amount += cgst_amount
        total_sgst_amount += sgst_amount
        total_igst_amount += igst_amount
        # Alignments for cells
        for idx, cell in enumerate(row_cells):
            if idx in [2, 4, 6, 8]:  # Right align these cells
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:  # Center align others
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    summary_row_cells = table.add_row().cells
    summary_row_cells[0].text = "Total"
    summary_row_cells[1].text = " "  # Empty HSN/SAC cell
    summary_row_cells[2].text = " "  # Total of Taxable Value
    summary_row_cells[3].text = " "  # Empty CGST Percentage cell
    summary_row_cells[4].text = " "  # Total of CGST Amount
    summary_row_cells[5].text = " "  # Empty SGST Percentage cell
    summary_row_cells[6].text = " "  # Total of SGST Amount
    summary_row_cells[7].text = " "  # Empty IGST Percentage cell
    summary_row_cells[8].text = " "  # Total of IGST Amount

    # Set background color for the summary row
    for cell in summary_row_cells:
        cell._element.get_or_add_tcPr().append(
            parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        )

        # Align text to the right for the appropriate cells (all cells except the first one)
        if cell != summary_row_cells[1]:  # Skip HSN/SAC cell
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            # HSN/SAC cell should be centered
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Align the first "Total" text to the center
    summary_row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Bold the amount fields (total_taxable_value, total_cgst_amount, total_sgst_amount, total_igst_amount)
    # Specifically bold the cells that contain the totals (Taxable Value, CGST Amount, SGST Amount, IGST Amount)
    total_taxable_value_cell = summary_row_cells[2]
    total_taxable_value_paragraph = total_taxable_value_cell.paragraphs[0]
    run = total_taxable_value_paragraph.add_run(str(total_taxable_value))
    run.bold = True

    total_cgst_amount_paragraph = summary_row_cells[4].paragraphs[0]
    run = total_cgst_amount_paragraph.add_run(f"{total_cgst_amount:.2f}")
    run.bold = True

    total_sgst_amount_paragraph = summary_row_cells[6].paragraphs[0]
    run = total_sgst_amount_paragraph.add_run(f"{total_sgst_amount:.2f}")
    run.bold = True

    total_igst_amount_paragraph = summary_row_cells[8].paragraphs[0]
    run = total_igst_amount_paragraph.add_run(f"{total_igst_amount:.2f}")
    run.bold = True

    return table  # Return the table with all the updates


# function generate terms and condition
def generate_terms_condition(table, datalist):
    terms_conditions_html = datalist["items"][0]["termsConditionsText"]
    terms_conditions_plain = terms_conditions_html.replace('<ol>', '').replace('</ol>', '')  # Remove <ol> tags
    terms_conditions_plain = terms_conditions_plain.replace('<li>', '• ').replace('</li>','\n')  # Add bullet points and newline
    for row in table.rows:
        for cell in row.cells:
            if "{termsandcondition}" in cell.text:
                # Find the paragraph that contains {termsandcondition}
                for p in cell.paragraphs:
                    if "{termsandcondition}" in p.text:
                        p.clear()  # Remove all text in the paragraph
                        p.add_run(terms_conditions_plain.strip())
                        p.alignment = p.alignment


def generate_document(json_data, Isquotation=False, status=None):
    datalist = json_data
    current_os = platform.system().lower()
    if current_os == 'windows':
        if Isquotation and status == "Draft":
            temlocations = BASE_DIR + "\static\PDF_TEMP\Final Quotation Print Format Draft.docx"
        elif Isquotation and status != "Draft":
            temlocations = BASE_DIR + "\static\PDF_TEMP\Final Quotation Print Format.docx"
        elif not Isquotation and status == "Draft":
            temlocations = BASE_DIR + "\static\PDF_TEMP\SalesOrderTemplate Draft.docx"
        elif not Isquotation and status != "Draft":
            temlocations = BASE_DIR + "\static\PDF_TEMP\SalesOrderTemplate.docx"
    else:
        if Isquotation and status == "Draft":
            temlocations = BASE_DIR + "/static/PDF_TEMP/Final Quotation Print Format Draft.docx"
        elif Isquotation and status != "Draft":
            temlocations = BASE_DIR + "/static/PDF_TEMP/Final Quotation Print Format.docx"
        elif not Isquotation and status == "Draft":
            temlocations = BASE_DIR + "/static/PDF_TEMP/SalesOrderTemplate Draft.docx"
        elif not Isquotation and status != "Draft":
            temlocations = BASE_DIR + "/static/PDF_TEMP/SalesOrderTemplate.docx"

    doc = Document(temlocations)

    for index, table in enumerate(doc.tables):
        if index == 0:
            if Isquotation:
                generate_quotation(table, datalist)
            else:
                generate_salesorder(table, datalist)
        if index == 1:
            final_amount = generate_item_detail(table, datalist)
        if index == 3:
            generate_amount(table, datalist, final_amount)
        if index == 4:
            generate_amount_in_words(table, datalist)
        if index == 5:
            generate_tax_data(table, datalist)
        if index == 6:
            generate_terms_condition(table, datalist)

    for table in doc.tables:
        set_cell_font_style(table, Isquotation)
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
        modified_document_path = temp_docx.name
        doc.save(modified_document_path)

    # Convert DOCX to PDF
    pdf_path = None
    try:
        pdf_path = convert_docx_to_pdf(modified_document_path)
    except Exception as e:
        print(e, "----->")

    return pdf_path


import shutil
import tempfile


def convert_docx_to_pdf(doc_path):
    try:
        if not doc_path:
            raise ValueError("The document path cannot be None or empty.")

        # Check if the input DOC/DOCX file exists
        if not os.path.isfile(doc_path):
            raise FileNotFoundError(f"The file {doc_path} does not exist.")
        # Get the directory of the input file and create a temporary directory for the PDF output
        output_dir = os.path.dirname(doc_path)
        temp_pdf_dir = tempfile.mkdtemp()  # Create a temporary directory to store the PDF

        # Define the path for the converted PDF in the temp directory
        temp_pdf_path = os.path.join(temp_pdf_dir, "tmp_output.pdf")

        # Get the system's OS and set the path for LibreOffice accordingly
        current_os = platform.system().lower()
        if current_os == 'windows':
            libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
            if not os.path.isfile(libreoffice_path):
                raise FileNotFoundError(f"LibreOffice executable not found at {libreoffice_path}")
        elif current_os == 'linux':
            libreoffice_path = "libreoffice"
        else:
            raise OSError(f"Unsupported OS: {current_os}. Only Windows and Linux are supported.")
     
        # Run LibreOffice in headless mode to convert DOC/DOCX to PDF
        result = subprocess.run(
            [libreoffice_path, '--headless', '--convert-to', 'pdf', '--outdir', temp_pdf_dir, doc_path],
            check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
        )
        # result = subprocess.run(
        #     [libreoffice_path, '--headless', '--convert-to', 'pdf:writer_pdf_Export', '--outdir', temp_pdf_dir, doc_path],
        #     check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
        # )
    #     result = subprocess.run(
    #     [
    #         libreoffice_path,
    #         '--headless',
    #         '--convert-to',
    #         'pdf:writer_pdf_Export',
    #         '--outdir',
    #         temp_pdf_dir,
    #         doc_path
    #     ],
    #     check=True,
    #     stdout=subprocess.PIPE,
    #     stderr=subprocess.PIPE,
    #     text=True
    # )

        # Capture any output or error for debugging
        if result.stdout:
            print("LibreOffice Output:", result.stdout)
        if result.stderr:
            print("LibreOffice Error:", result.stderr)

        # The generated PDF should be in the temporary directory now, rename it to the desired name
        # Assuming LibreOffice generates a PDF with a temporary name like tmp1234.pdf
        for file in os.listdir(temp_pdf_dir):
            if file.endswith(".pdf"):
                temp_pdf_path = os.path.join(temp_pdf_dir, file)
                break  # Stop after finding the first PDF file

        # Rename the PDF to the desired output name (e.g., output.pdf)
        final_pdf_path = doc_path.replace('.docx', '.pdf').replace('.doc', '.pdf')  # Output in same directory
        shutil.move(temp_pdf_path, final_pdf_path)
 
        # Check if the final PDF exists at the expected location
        if os.path.exists(final_pdf_path):

            return final_pdf_path
        else:
            raise Exception(
                f"DOC/DOCX to PDF conversion failed. PDF file was not created. Expected at {final_pdf_path}")

    except Exception as e:

        raise




ONLYOFFICE_CONTAINER_NAME = "onlyoffice-document-server"

def ensure_onlyoffice_running():
    try:
        # Check if container exists
        result = subprocess.run(
            ["docker", "ps", "-a", "--filter", f"name={ONLYOFFICE_CONTAINER_NAME}", "--format", "{{.Names}}"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        containers = result.stdout.strip().splitlines()

        if ONLYOFFICE_CONTAINER_NAME in containers:
            # Start if not already running
            running = subprocess.run(
                ["docker", "inspect", "-f", "{{.State.Running}}", ONLYOFFICE_CONTAINER_NAME],
                stdout=subprocess.PIPE,
                text=True,
            )
            if running.stdout.strip() != "true":
                subprocess.run(["docker", "start", ONLYOFFICE_CONTAINER_NAME], check=True)
                print("✅ Started existing OnlyOffice container.")
            else:
                print("✅ OnlyOffice container is already running.")
        else:
            # Run a new container
            subprocess.run([
                "docker", "run", "-i", "-t", "-d",
                "-p", "8000:80",
                "--name", ONLYOFFICE_CONTAINER_NAME,
                "onlyoffice/documentserver"
            ], check=True)
            print("✅ Launched new OnlyOffice container.")
    except Exception as e:
        raise RuntimeError(f"Failed to start OnlyOffice container: {e}")

def convert_docx_to_pdf_cross_platform(docx_path):
    if not docx_path or not os.path.isfile(docx_path):
        raise FileNotFoundError(f"File not found: {docx_path}")

    system_os = platform.system().lower()
    pdf_path = docx_path.replace('.docx', '.pdf').replace('.doc', '.pdf')

    if system_os == "windows":
        try:
            import comtypes.client
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(docx_path)
            doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
            doc.Close()
            word.Quit()
            if not os.path.exists(pdf_path):
                raise Exception("PDF not created on Windows.")
            return pdf_path
        except Exception as e:
            raise RuntimeError(f"Windows conversion failed: {e}")

    elif system_os == "linux":
        try:
            ensure_onlyoffice_running()  # Ensure OnlyOffice is running before converting

            with open(docx_path, "rb") as f:
                docx_base64 = base64.b64encode(f.read()).decode()

            payload = {
                "async": False,
                "filetype": "docx",
                "outputtype": "pdf",
                "title": "ConvertedFile",
                "key": os.path.basename(docx_path),
                "file": docx_base64
            }

            response = requests.post(
                "http://localhost:8000/ConvertService.ashx",
                headers={"Content-Type": "application/json"},
                data=json.dumps(payload),
                timeout=30
            )
            result = response.json()
            file_url = result.get("fileUrl")
            if not file_url:
                raise Exception(f"OnlyOffice conversion failed: {result}")

            # Download and save the final PDF
            pdf_response = requests.get(file_url)
            with open(pdf_path, "wb") as f:
                f.write(pdf_response.content)

            return pdf_path

        except Exception as e:
            raise RuntimeError(f"OnlyOffice conversion failed on Linux: {e}")

    else:
        raise OSError(f"Unsupported OS: {system_os}")
